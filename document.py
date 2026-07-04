from docx import Document

def generate_doc(text):

    doc = Document()

    doc.add_heading("AI Generated Document", level=1)

    doc.add_paragraph(text)

    path = "outputs/report.docx"

    doc.save(path)

    return path